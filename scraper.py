#first, let's import our package of word document to read it, and re to use regular expressions
import re
import docx
import openpyxl
#now, let's create our spreadsheet that the data will be populated to
wb = openpyxl.Workbook()
sheet = wb.get_sheet_by_name('Sheet')
#and now we must make the ultimate sheet value for increasing the rows one by one
sheet_row = 1
#now, let's open and save the document itself
doc = docx.Document('Test3.docx')
#let's read how many lines (paragraphs) will be in this document
n = len(doc.paragraphs)
#the first line of every document is the region
#because this will stay constant for the entire document, we can define it here
region = doc.paragraphs[0].text
print(region)
hoard_name = ""
tpq = ""
dynasty = ""
date = ""
weight = ""
comment = "fragment"
leader = ""
mint = ""
type = ""
#now, let's define our various functions for finding different variables
#first, we'll do hoard name
def hoard_name_match(strg, search=re.compile('^[\d]+[.]\s+[A-Z]+').search):
	return bool(search(strg))
	
#now, let's try for type
def type_match(strg, search=re.compile('([A-C])[.]\s[A-Z\s]+\s\d+').search):
	return bool(search(strg))
	
#and for dynasty
def dynasty_match(strg, search=re.compile('([IVXCL]+)[.]\s[A-Za-z\s]+\s\d+').search):
	return bool(search(strg))
	
#and now for the mint date and weight with comments line
def mint_line_match1(strg, search=re.compile('^[\d]+\s[A-Za][a-z]+').search):
	return bool(search(strg))
	
def mint_line_match2(strg, search=re.compile('^[\d]+\s[A-Za][a-z]+[-]+').search):
	return bool(search(strg))
	
#and the leader
def leader_match1(strg, search=re.compile('^[al\W]+[A-Za-z\s]+\s+\d+').search):
	return bool(search(strg))
	
def leader_match2(strg, search=re.compile('^[A-Za-z\s]+\s+\d+').search):
	return bool(search(strg))
	
def mint_match1(strg, search=re.compile('^[A-Za-z]+[-][A-Za-z]+').search):
	return bool(search(strg))
	
def mint_match2(strg, search=re.compile('^[A-Za-z]').search):
	return bool(search(strg))
	
def date_match1(strg, search=re.compile('^[\d]+/[\d]+').search):
	return bool(search(strg))
	
def date_match2(strg, search=re.compile('^[\d/]+[-][\d/]+').search):
	return bool(search(strg))
	
def date_match3(strg, search=re.compile('^[\d][\d][\d]').search):
	return bool(search(strg))
	
def date_match4(strg, search=re.compile('^[\d][\d][\d][\d]').search):
	return bool(search(strg))
	
def weight_match(strg, search=re.compile('^[\d][.][\dg]+').search):
	return bool(search(strg))
	
def just_multiple(strg, search=re.compile('[\d]+[&]').search):
	return bool(search(strg))
	
def frag_multiple(strg, search=re.compile('[\d]+[$][&]').search):
	return bool(search(strg))
	
def just_frag(strg, search=re.compile('^[#]').search):
	return bool(search(strg))
	
def just_more1(strg, search=re.compile('^[\d]\Z').search):
	return bool(search(strg))\
	
def just_more2(strg, search=re.compile('^[\d]\Z').search):
	return bool(search(strg))\
		
#now let's make our for loop, looping n times where n is the number of lines in the document
for i in range(n):
	#this will create an easier variable to work with
	test_text = doc.paragraphs[i].text
	#this splits up the string pre-emptively by space
	test_text_split = test_text.split()
	#this tests for the length of the line
	y = len(test_text_split)
	#this checks if the line matches the hoard name case
	if i-1 >= 0: 
		if type_match(test_text):
			type = str(test_text_split[1:(y-1)])
		elif hoard_name_match(doc.paragraphs[i-1].text):
			tpqlen = len(doc.paragraphs[i].runs)
			for x3 in range(tpqlen):
				if doc.paragraphs[i].runs[x3].bold == True:
					tpq = doc.paragraphs[i].runs[x3].text
		elif leader_match1(test_text) or leader_match2(test_text):
			leader = str(test_text_split[0:(y-1)])
		elif dynasty_match(test_text):
			dynasty = str(test_text_split[1:(y-1)])
			#also need to reset leader field with each new dynasty, as that dynasty may not have a leader type
			leader = ""
		elif hoard_name_match(test_text):
			hoard_name = str(test_text_split[1:y])
			weight = ""
			tpq = ""
		elif mint_line_match1(test_text) or mint_line_match2(test_text):
			mint = ""
			mint_line = test_text
			y = len(test_text_split)
			for w in range(y):
				if mint_match1(test_text_split[w]) or mint_match2(test_text_split[w]):
					if mint == "":
						mint = test_text_split[w]
					else:
						mint = mint + " " + test_text_split[w]
				if date_match1(test_text_split[w]) or date_match2(test_text_split[w]) or date_match3(test_text_split[w]) or date_match4(test_text_split[w]):
					date = test_text_split[w]
					if w == y-1:
						print(region, hoard_name, tpq, type, dynasty, leader, mint, date)
						sheet.cell(row=sheet_row, column=10).value = region
						sheet.cell(row=sheet_row, column=1).value = hoard_name
						sheet.cell(row=sheet_row, column=2).value = tpq
						sheet.cell(row=sheet_row, column=3).value = str(type)
						sheet.cell(row=sheet_row, column=4).value = dynasty
						sheet.cell(row=sheet_row, column=5).value = leader
						sheet.cell(row=sheet_row, column=6).value = mint
						sheet.cell(row=sheet_row, column=7).value = date
						sheet_row = sheet_row + 1
					if w+2 < y:
						if weight_match(test_text_split[w+1]) and just_frag(test_text_split[w+2]):
							weight = test_text_split[w+1]
							comment = "fragment"
							print(region, hoard_name, tpq, type, dynasty, leader, mint, date, weight, comment)
							sheet.cell(row=sheet_row, column=10).value = region
							sheet.cell(row=sheet_row, column=1).value = hoard_name
							sheet.cell(row=sheet_row, column=2).value = tpq
							sheet.cell(row=sheet_row, column=3).value = str(type)
							sheet.cell(row=sheet_row, column=4).value = dynasty
							sheet.cell(row=sheet_row, column=5).value = leader
							sheet.cell(row=sheet_row, column=6).value = mint
							sheet.cell(row=sheet_row, column=7).value = date
							sheet.cell(row=sheet_row, column=8).value = weight
							sheet.cell(row=sheet_row, column=9).value = comment
							sheet_row = sheet_row + 1
						elif weight_match(test_text_split[w+1]):
							weight = test_text_split[w+1]
							print(region, hoard_name, tpq, type, dynasty, leader, mint, date, weight)
							sheet.cell(row=sheet_row, column=10).value = region
							sheet.cell(row=sheet_row, column=1).value = hoard_name
							sheet.cell(row=sheet_row, column=2).value = tpq
							sheet.cell(row=sheet_row, column=3).value = str(type)
							sheet.cell(row=sheet_row, column=4).value = dynasty
							sheet.cell(row=sheet_row, column=5).value = leader
							sheet.cell(row=sheet_row, column=6).value = mint
							sheet.cell(row=sheet_row, column=7).value = date
							sheet.cell(row=sheet_row, column=8).value = weight
							sheet_row = sheet_row + 1
						elif just_multiple(test_text_split[w+1]):
							j = int(filter(unicode.isdigit, test_text_split[w+1]))
							z = 0
							h = 0
							while z < j:
								if w+3+h < y:
									if weight_match(test_text_split[w+2+h]) and just_frag(test_text_split[w+3+h]):
										weight = test_text_split[w+2+h]
										comment = "fragment"
										print(region, hoard_name, tpq, type, dynasty, leader, mint, date, weight, comment)
										sheet.cell(row=sheet_row, column=10).value = region
										sheet.cell(row=sheet_row, column=1).value = hoard_name
										sheet.cell(row=sheet_row, column=2).value = tpq
										sheet.cell(row=sheet_row, column=3).value = type
										sheet.cell(row=sheet_row, column=4).value = dynasty
										sheet.cell(row=sheet_row, column=5).value = leader
										sheet.cell(row=sheet_row, column=6).value = mint
										sheet.cell(row=sheet_row, column=7).value = date
										sheet.cell(row=sheet_row, column=8).value = weight
										sheet.cell(row=sheet_row, column=9).value = comment
										sheet_row = sheet_row + 1
										z = z + 1
										h = h + 1
									elif weight_match(test_text_split[w+2+h]):
										weight = test_text_split[w+2+h]
										print(region, hoard_name, tpq, type, dynasty, leader, mint, date, weight)
										sheet.cell(row=sheet_row, column=10).value = region
										sheet.cell(row=sheet_row, column=1).value = hoard_name
										sheet.cell(row=sheet_row, column=2).value = tpq
										sheet.cell(row=sheet_row, column=3).value = type
										sheet.cell(row=sheet_row, column=4).value = dynasty
										sheet.cell(row=sheet_row, column=5).value = leader
										sheet.cell(row=sheet_row, column=6).value = mint
										sheet.cell(row=sheet_row, column=7).value = date
										sheet.cell(row=sheet_row, column=8).value = weight
										sheet_row = sheet_row + 1
										z = z + 1
										h = h + 1
									elif date_match1(test_text_split[w+2]) or date_match2(test_text_split[w+2]) or date_match3(test_text_split[w+2]) or date_match4(test_text_split[w+2]):
										print(region, hoard_name, tpq, type, dynasty, leader, mint, date)
										sheet.cell(row=sheet_row, column=10).value = region
										sheet.cell(row=sheet_row, column=1).value = hoard_name
										sheet.cell(row=sheet_row, column=2).value = tpq
										sheet.cell(row=sheet_row, column=3).value = type
										sheet.cell(row=sheet_row, column=4).value = dynasty
										sheet.cell(row=sheet_row, column=5).value = leader
										sheet.cell(row=sheet_row, column=6).value = mint
										sheet.cell(row=sheet_row, column=7).value = date
										sheet_row = sheet_row + 1
										z = z + 1
										h = h + 1
									else:
										h = h + 1
								elif w+2+h < y:
									if weight_match(test_text_split[w+2+h]):
										weight = test_text_split[w+2+h]
										print(region, hoard_name, tpq, type, dynasty, leader, mint, date, weight)
										sheet.cell(row=sheet_row, column=10).value = region
										sheet.cell(row=sheet_row, column=1).value = hoard_name
										sheet.cell(row=sheet_row, column=2).value = tpq
										sheet.cell(row=sheet_row, column=3).value = type
										sheet.cell(row=sheet_row, column=4).value = dynasty
										sheet.cell(row=sheet_row, column=5).value = leader
										sheet.cell(row=sheet_row, column=6).value = mint
										sheet.cell(row=sheet_row, column=7).value = date
										sheet.cell(row=sheet_row, column=8).value = weight
										sheet_row = sheet_row + 1
										z = z + 1
										h = h + 1
									elif date_match1(test_text_split[w+2]) or date_match2(test_text_split[w+2]) or date_match3(test_text_split[w+2]) or date_match4(test_text_split[w+2]):
										print(region, hoard_name, tpq, type, dynasty, leader, mint, date)
										sheet.cell(row=sheet_row, column=10).value = region
										sheet.cell(row=sheet_row, column=1).value = hoard_name
										sheet.cell(row=sheet_row, column=2).value = tpq
										sheet.cell(row=sheet_row, column=3).value = type
										sheet.cell(row=sheet_row, column=4).value = dynasty
										sheet.cell(row=sheet_row, column=5).value = leader
										sheet.cell(row=sheet_row, column=6).value = mint
										sheet.cell(row=sheet_row, column=7).value = date
										sheet_row = sheet_row + 1
										z = z + 1
										h = h + 1
									else:
										h = h + 1
								else:
									h = h + 1
									z = z + 1
						elif frag_multiple(test_text_split[w+1]):
							j = int(filter(unicode.isdigit, test_text_split[w+1]))
							z = 0
							h = 0
							while z < j:
								if w+2+h < y:
									if weight_match(test_text_split[w+2+h]):
										weight = test_text_split[w+2+h]
										comment = "fragment"
										print(region, hoard_name, tpq, type, dynasty, leader, mint, date, weight, comment)
										sheet.cell(row=sheet_row, column=10).value = region
										sheet.cell(row=sheet_row, column=1).value = hoard_name
										sheet.cell(row=sheet_row, column=2).value = tpq
										sheet.cell(row=sheet_row, column=3).value = type
										sheet.cell(row=sheet_row, column=4).value = dynasty
										sheet.cell(row=sheet_row, column=5).value = leader
										sheet.cell(row=sheet_row, column=6).value = mint
										sheet.cell(row=sheet_row, column=7).value = date
										sheet.cell(row=sheet_row, column=8).value = weight
										sheet.cell(row=sheet_row, column=9).value = comment
										sheet_row = sheet_row + 1
										z = z + 1
										h = h + 1
									else:
										h = h + 1
								else:
									z = z + 1
						elif date_match1(test_text_split[w+1]) or date_match2(test_text_split[w+1]) or date_match3(test_text_split[w+1]):
							print(region, hoard_name, tpq, type, dynasty, leader, mint, date)
							sheet.cell(row=sheet_row, column=10).value = region
							sheet.cell(row=sheet_row, column=1).value = hoard_name
							sheet.cell(row=sheet_row, column=2).value = tpq
							sheet.cell(row=sheet_row, column=3).value = type
							sheet.cell(row=sheet_row, column=4).value = dynasty
							sheet.cell(row=sheet_row, column=5).value = leader
							sheet.cell(row=sheet_row, column=6).value = mint
							sheet.cell(row=sheet_row, column=7).value = date
							sheet_row = sheet_row + 1			
						elif just_frag(test_text_split[w+1]):
							print(region, hoard_name, tpq, type, dynasty, leader, mint, date)
							sheet.cell(row=sheet_row, column=10).value = region
							sheet.cell(row=sheet_row, column=1).value = hoard_name
							sheet.cell(row=sheet_row, column=2).value = tpq
							sheet.cell(row=sheet_row, column=3).value = type
							sheet.cell(row=sheet_row, column=4).value = dynasty
							sheet.cell(row=sheet_row, column=5).value = leader
							sheet.cell(row=sheet_row, column=6).value = mint
							sheet.cell(row=sheet_row, column=7).value = date
							sheet.cell(row=sheet_row, column=9).value = comment
							sheet_row = sheet_row + 1
						elif just_more1(test_text_split[w+1]) or just_more2(test_text_split[w+1]):
							j = int(filter(unicode.isdigit, test_text_split[w+1]))
							z = 0
							h = 0
							while z < j:
								print(region, hoard_name, tpq, type, dynasty, leader, mint, date)
								sheet.cell(row=sheet_row, column=10).value = region
								sheet.cell(row=sheet_row, column=1).value = hoard_name
								sheet.cell(row=sheet_row, column=2).value = tpq
								sheet.cell(row=sheet_row, column=3).value = str(type)
								sheet.cell(row=sheet_row, column=4).value = dynasty
								sheet.cell(row=sheet_row, column=5).value = leader
								sheet.cell(row=sheet_row, column=6).value = mint
								sheet.cell(row=sheet_row, column=7).value = date
								sheet_row = sheet_row + 1
								z = z + 1
								h = h + 1
						else:
							print(region, hoard_name, tpq, type, dynasty, leader, mint, date)
							sheet.cell(row=sheet_row, column=10).value = region
							sheet.cell(row=sheet_row, column=1).value = hoard_name
							sheet.cell(row=sheet_row, column=2).value = tpq
							sheet.cell(row=sheet_row, column=3).value = str(type)
							sheet.cell(row=sheet_row, column=4).value = dynasty
							sheet.cell(row=sheet_row, column=5).value = leader
							sheet.cell(row=sheet_row, column=6).value = mint
							sheet.cell(row=sheet_row, column=7).value = date
							sheet_row = sheet_row + 1
					elif w + 1 < y:
						if weight_match(test_text_split[w+1]):
							weight = test_text_split[w+1]
							print(region, hoard_name, tpq, type, dynasty, leader, mint, date, weight)
							sheet.cell(row=sheet_row, column=10).value = region
							sheet.cell(row=sheet_row, column=1).value = hoard_name
							sheet.cell(row=sheet_row, column=2).value = tpq
							sheet.cell(row=sheet_row, column=3).value = str(type)
							sheet.cell(row=sheet_row, column=4).value = dynasty
							sheet.cell(row=sheet_row, column=5).value = leader
							sheet.cell(row=sheet_row, column=6).value = mint
							sheet.cell(row=sheet_row, column=7).value = date
							sheet.cell(row=sheet_row, column=8).value = weight
							sheet_row = sheet_row + 1
						elif just_multiple(test_text_split[w+1]):
							j = int(filter(unicode.isdigit, test_text_split[w+1]))
							z = 0
							h = 0
							while z < j:
								if w+3+h < y:
									if weight_match(test_text_split[w+2+h]) and just_frag(test_text_split[w+3+h]):
										weight = test_text_split[w+2+h]
										comment = "fragment"
										print(region, hoard_name, tpq, type, dynasty, leader, mint, date, weight, comment)
										sheet.cell(row=sheet_row, column=10).value = region
										sheet.cell(row=sheet_row, column=1).value = hoard_name
										sheet.cell(row=sheet_row, column=2).value = tpq
										sheet.cell(row=sheet_row, column=3).value = type
										sheet.cell(row=sheet_row, column=4).value = dynasty
										sheet.cell(row=sheet_row, column=5).value = leader
										sheet.cell(row=sheet_row, column=6).value = mint
										sheet.cell(row=sheet_row, column=7).value = date
										sheet.cell(row=sheet_row, column=8).value = weight
										sheet.cell(row=sheet_row, column=9).value = comment
										sheet_row = sheet_row + 1
										z = z + 1
										h = h + 1
									elif weight_match(test_text_split[w+2+h]):
										weight = test_text_split[w+2+h]
										print(region, hoard_name, tpq, type, dynasty, leader, mint, date, weight)
										sheet.cell(row=sheet_row, column=10).value = region
										sheet.cell(row=sheet_row, column=1).value = hoard_name
										sheet.cell(row=sheet_row, column=2).value = tpq
										sheet.cell(row=sheet_row, column=3).value = type
										sheet.cell(row=sheet_row, column=4).value = dynasty
										sheet.cell(row=sheet_row, column=5).value = leader
										sheet.cell(row=sheet_row, column=6).value = mint
										sheet.cell(row=sheet_row, column=7).value = date
										sheet.cell(row=sheet_row, column=8).value = weight
										sheet_row = sheet_row + 1
										z = z + 1
										h = h + 1
									else:
										h = h + 1
								elif w+2+h < y:
									if weight_match(test_text_split[w+2+h]):
										weight = test_text_split[w+2+h]
										print(region, hoard_name, tpq, type, dynasty, leader, mint, date, weight)
										sheet.cell(row=sheet_row, column=10).value = region
										sheet.cell(row=sheet_row, column=1).value = hoard_name
										sheet.cell(row=sheet_row, column=2).value = tpq
										sheet.cell(row=sheet_row, column=3).value = type
										sheet.cell(row=sheet_row, column=4).value = dynasty
										sheet.cell(row=sheet_row, column=5).value = leader
										sheet.cell(row=sheet_row, column=6).value = mint
										sheet.cell(row=sheet_row, column=7).value = date
										sheet.cell(row=sheet_row, column=8).value = weight
										sheet_row = sheet_row + 1
										z = z + 1
										h = h + 1
									else:
										h = h + 1
								else:
									h = h + 1
									z = z + 1
						elif frag_multiple(test_text_split[w+1]):
							j = int(filter(str.isdigit, test_text_split[w+1]))
							z = 0
							h = 0
							while z < j:
								if weight_match(test_text_split[w+2+h]):
									weight = test_text_split[w+2+h]
									comment = "fragment"
									print(region, hoard_name, tpq, type, dynasty, leader, mint, date, weight, comment)
									sheet.cell(row=sheet_row, column=10).value = region
									sheet.cell(row=sheet_row, column=1).value = hoard_name
									sheet.cell(row=sheet_row, column=2).value = tpq
									sheet.cell(row=sheet_row, column=3).value = type
									sheet.cell(row=sheet_row, column=4).value = dynasty
									sheet.cell(row=sheet_row, column=5).value = leader
									sheet.cell(row=sheet_row, column=6).value = mint
									sheet.cell(row=sheet_row, column=7).value = date
									sheet.cell(row=sheet_row, column=8).value = weight
									sheet.cell(row=sheet_row, column=9).value = comment
									sheet_row = sheet_row + 1
									z = z + 1
									h = h + 1
								else:
									h = h + 1
						elif date_match1(test_text_split[w+1]) or date_match2(test_text_split[w+1]) or date_match3(test_text_split[w+1]):
							print(region, hoard_name, tpq, type, dynasty, leader, mint, date)
							sheet.cell(row=sheet_row, column=10).value = region
							sheet.cell(row=sheet_row, column=1).value = hoard_name
							sheet.cell(row=sheet_row, column=2).value = tpq
							sheet.cell(row=sheet_row, column=3).value = type
							sheet.cell(row=sheet_row, column=4).value = dynasty
							sheet.cell(row=sheet_row, column=5).value = leader
							sheet.cell(row=sheet_row, column=6).value = mint
							sheet.cell(row=sheet_row, column=7).value = date
							sheet_row = sheet_row + 1
						elif just_frag(test_text_split[w+1]):
							comment = "fragment"
							print(region, hoard_name, tpq, type, dynasty, leader, mint, date, comment)
							sheet.cell(row=sheet_row, column=10).value = region
							sheet.cell(row=sheet_row, column=1).value = hoard_name
							sheet.cell(row=sheet_row, column=2).value = tpq
							sheet.cell(row=sheet_row, column=3).value = type
							sheet.cell(row=sheet_row, column=4).value = dynasty
							sheet.cell(row=sheet_row, column=5).value = leader
							sheet.cell(row=sheet_row, column=6).value = mint
							sheet.cell(row=sheet_row, column=7).value = date
							sheet.cell(row=sheet_row, column=9).value = comment
							sheet_row = sheet_row + 1
						elif just_more1(test_text_split[w+1]) or just_more2(test_text_split[w+1]):
							j = int(filter(unicode.isdigit, test_text_split[w+1]))
							z = 0
							h = 0
							while z < j:
								print(region, hoard_name, tpq, type, dynasty, leader, mint, date)
								sheet.cell(row=sheet_row, column=10).value = region
								sheet.cell(row=sheet_row, column=1).value = hoard_name
								sheet.cell(row=sheet_row, column=2).value = tpq
								sheet.cell(row=sheet_row, column=3).value = type
								sheet.cell(row=sheet_row, column=4).value = dynasty
								sheet.cell(row=sheet_row, column=5).value = leader
								sheet.cell(row=sheet_row, column=6).value = mint
								sheet.cell(row=sheet_row, column=7).value = date
								sheet_row = sheet_row + 1
								z = z + 1
								h = h + 1
						elif date_match1(test_text_split[w]) or date_match2(test_text_split[w]) or date_match3(test_text_split[w]):
								date = test_text_split[w]
								print(region, hoard_name, tpq, type, dynasty, leader, mint, date)
								sheet.cell(row=sheet_row, column=10).value = region
								sheet.cell(row=sheet_row, column=1).value = hoard_name
								sheet.cell(row=sheet_row, column=2).value = tpq
								sheet.cell(row=sheet_row, column=3).value = type
								sheet.cell(row=sheet_row, column=4).value = dynasty
								sheet.cell(row=sheet_row, column=5).value = leader
								sheet.cell(row=sheet_row, column=6).value = mint
								sheet.cell(row=sheet_row, column=7).value = date
								sheet_row = sheet_row + 1
wb.save('Test3.xlsx')